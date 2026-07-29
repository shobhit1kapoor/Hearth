from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission"
OUT.mkdir(exist_ok=True)

PINE = "173F32"
PINE_2 = "285C49"
INK = "20302A"
MUTED = "5F6D67"
AMBER = "8A5B17"
PALE = "F2F4F1"
MINT = "E6F0EB"
RED = "8E3027"
WHITE = "FFFFFF"
LINE = "D8DDD9"


def set_run(run, size=11, bold=False, italic=False, color=INK, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_field(paragraph, total=None):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    if total:
        set_run(paragraph.add_run(f" of {total}"), size=9, color=MUTED)


def configure_document(doc: Document, running_label: str, total_pages=None):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for name, size, color, before, after in (
        ("Heading 1", 15, PINE, 12, 7),
        ("Heading 2", 12.5, PINE_2, 8, 5),
        ("Heading 3", 11.5, INK, 6, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208
    header = sec.header.paragraphs[0]
    header.text = ""
    set_run(header.add_run(running_label), size=9, bold=True, color=MUTED)
    header.paragraph_format.space_after = Pt(0)
    footer = sec.footer.paragraphs[0]
    add_page_field(footer, total_pages)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(text.upper()), size=9, bold=True, color=AMBER)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def add_bullets(doc, items: Iterable[str], compact=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if compact:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.08
        set_run(p.add_run(item))


def add_table(doc, headers, rows, widths=None, font_size=9):
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PALE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(str(label)), size=font_size, bold=True, color=PINE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(table, widths)
    return table


def add_callout(doc, title, text, color=MINT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(title + " "), bold=True, color=PINE)
    set_run(p.add_run(text))
    set_table_geometry(table, [9360])
    return table


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), PINE_2)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


APPLICATION_PAGES = [
    {
        "section": "Section 1: Understanding of Need and Solution Design",
        "title": "The operating gap after discharge",
        "kicker": "Narrative page 1 of 15",
        "paragraphs": [
            "HEARTH addresses a practical failure that appears after a person leaves a hospital or moves through another care transition: information exists, yet the work required to carry it out is scattered across discharge papers, medication lists, portal messages, calls, calendars, text threads, and memory. Family caregivers become the unofficial integration layer. They must identify what must happen, decide who can do it, discover missing skill or equipment, protect the care recipient’s preferences, and verify that outside organizations actually responded.",
            "Most tools store information or generate reminders. HEARTH instead treats each obligation as an auditable care commitment object with an owner, source, safe window, dependencies, permission rule, completion evidence, backup, escalation path, safety level, and lifecycle state. The goal is not to replace a caregiver or clinician. It is to make hidden coordination work visible and executable while preserving the authority of the person receiving care.",
            "The Phase 1 prototype demonstrates this operating model in one synthetic transition-to-home case. Ten synthetic sources compile into 26 commitments. The initial mission is correctly marked NOT EXECUTABLE because high-risk gaps remain: conflicting insulin instructions, no qualified wound-care owner, missing equipment, caregiver capacity deficit, unaccepted transportation, and an unidentified instruction. Unaffected work may continue, but the system never silently converts uncertainty into a clinical instruction.",
        ],
        "callout": ("Core proposition.", "HEARTH converts fragmented care information into source-grounded, permission-aware, closed-loop responsibilities—and stops when the evidence or authority is insufficient."),
    },
    {
        "title": "Who experiences the burden and why current workflows fail",
        "kicker": "Narrative page 2 of 15",
        "paragraphs": [
            "The primary user is an unpaid adult caregiver coordinating care across a household and multiple organizations. The care recipient remains the center of permission and preference decisions. Other users include family helpers, home-health professionals, pharmacists, clinic staff, aging-network navigators, and transition-of-care teams, but none receives blanket access or authority.",
            "The failure is operational rather than informational. A caregiver can possess every document and still lack an executable plan. A discharge instruction may omit an owner. An appointment may be scheduled without a ride. A requested dressing change may lack trained help or supplies. A message may be sent but never acknowledged. A newer instruction may conflict with an older list. Ordinary task software does not distinguish a draft from an accepted responsibility, a sent message from an outcome, or an uncertain phrase from a verified instruction.",
            "HEARTH defines burden broadly: time, interactions, corrections, uncertainty, emotional effort, skills, equipment, scheduling, and the privacy cost of involving more people. Its capacity shield compares estimated work against declared availability and exposes clustered loads without diagnosing burnout. Its minimum-disclosure control gives a ride helper logistics while withholding diagnosis, medicines, insurance, and caregiver-private notes.",
        ],
        "bullets": [
            "Care recipient: purpose-specific, revocable control and supported decision-making.",
            "Primary caregiver: compilation, approval, correction, delegation, and evidence review.",
            "Helper: accepted task and only the fields required for that purpose.",
            "Professional: clinical ambiguity and high-risk conflicts remain within licensed authority.",
        ],
    },
    {
        "title": "Solution design and responsible use of AI",
        "kicker": "Narrative page 3 of 15",
        "paragraphs": [
            "HEARTH uses AI only where interpretation is useful and surrounds it with deterministic controls. A compiler can extract candidate responsibilities, deadlines, entities, source locations, conflicts, dependencies, and uncertainties. The deterministic layer then applies the lifecycle graph, safety level, permission scope, staleness, completion-evidence requirement, and audit receipt. High-risk output is prepared for review; it is not automatically activated or sent.",
            "The key unit is the care commitment object. Unlike a free-form summary, it preserves why the work exists, who owns it, which source supports it, what success looks like, what can block it, which data may be shared, and where to escalate. The mission status is calculated from unresolved critical and high findings. This produces a visible difference between READY, READY WITH CONTROLS, HOLD, and NOT EXECUTABLE.",
            "The interface is deliberately caregiver-facing: Today’s mission, Reality check, Mission board, Care inbox, Compilation review, Medication safety, Appointments, Family care circle, Capacity shield, Permission vault, Accountability receipts, Trust and privacy, Evidence and validation, Timed burden study, and a guided reviewer demo. All people and records in Phase 1 are synthetic and every external response is labeled as a simulation.",
        ],
        "table": (["AI may", "AI may not"], [
            ["Extract and prepare candidate responsibilities with provenance", "Diagnose, prescribe, choose a conflicting dose, or invent missing clinical meaning"],
            ["Detect conflicts, staleness, missing owners, and capacity gaps", "Represent a sent message as a completed outcome"],
            ["Draft minimum-necessary requests and questions", "Send, purchase, schedule, or change a record without authorized approval"],
        ], [4680, 4680]),
    },
    {
        "title": "Caregiver input: current evidence and corrective plan",
        "kicker": "Narrative page 4 of 15",
        "paragraphs": [
            "A complete audit of the repository and the supplied build materials found no consented caregiver transcript, recording, survey, interview note, coded record, or quote. Therefore this application does not claim that real caregivers shaped HEARTH, does not assign fictional participant identifiers, and does not present synthetic demo language as an authentic quotation. The current design is a challenge-aligned hypothesis that requires caregiver validation.",
            "This missing evidence is a material weakness against the caregiver-input criterion, not a footnote. The repository contains a ready-to-use consent and data-minimization method, interview guide, codebook, anonymized participant register, authentic-quote gate, thematic-analysis template, disconfirming-feedback register, insight-to-design matrix, and decision log. They remain visibly empty where source evidence is absent.",
            "Phase 2 begins with a compensated caregiver advisory group of 6–10 people who can veto unsafe authority language, set privacy defaults, prioritize integrations, approve participant materials, and review interpreted findings. A first discovery sample of 5–8 caregivers will intentionally include varied digital comfort, relationships, work constraints, disability access needs, and care complexity. Both supporting and opposing evidence will be traced to design decisions.",
        ],
        "callout": ("Submission gate.", "Before upload, the applicant should add verified caregiver evidence or accept that this criterion is only addressed through a transparent co-implementation plan."),
    },
    {
        "section": "Section 2: Implementation Approach",
        "title": "Technical architecture and execution controls",
        "kicker": "Narrative page 5 of 15",
        "paragraphs": [
            "Phase 1 is a functioning browser proof of concept at TRL 3. A Next.js/React interface presents a deterministic TypeScript mission engine. Synthetic sources and fixtures exercise extraction fields, conflicts, permissions, capacity, lifecycle, receipts, and evidence views. Build, type, unit, rendered-HTML, controlled validation, holdout, and accessibility checks are reproducible from the repository.",
            "A production architecture would separate ingestion, document interpretation, commitment storage, consent and identity, workflow adapters, audit events, evaluation telemetry, and the caregiver interface. Source content is immutable; corrections append a new version. Every object carries a household identifier. Purpose-specific grants constrain actor, fields, action, expiry, and revocation. External actions use prepare/approve/send/acknowledge/outcome states rather than a single completion flag.",
            "Safe failure is the default. If the interpretation service is unavailable, the system retains the deterministic task view and suspends new interpretations. If a provider or pharmacy does not respond, the responsibility remains open and follows a configured escalation window. Prompt-injection patterns inside source text are quarantined as data rather than executed as instructions.",
        ],
        "table": (["Layer", "Phase 1", "Phase 2 gate"], [
            ["Experience", "Responsive browser prototype", "Caregiver-tested accessible workflows"],
            ["Mission engine", "Deterministic safety and state helpers", "Versioned service with policy tests"],
            ["AI interpretation", "Controlled fixtures and abstention", "Model card, monitoring, subgroup evaluation"],
            ["Integrations", "Labeled simulations", "Partner sandbox and governance approval"],
        ], [1800, 3300, 4260]),
    },
    {
        "title": "Safety, privacy, security, and human authority",
        "kicker": "Narrative page 6 of 15",
        "paragraphs": [
            "HEARTH uses five safety levels from H0 administrative support to H4 emergency boundary. H3 and H4 work requires a verified source and qualified human review. The exact Protocol 9-Delta test demonstrates the boundary: the system says it cannot identify the protocol, will not invent or apply a meaning, and requests an approved source or qualified review.",
            "Privacy is not a generic role label. A disclosure is tied to purpose and data categories. Cross-household requests, expired grants, revoked grants, and requests without a defined purpose are denied. A correction preserves the original text, corrected text, reason, actor, and time. Audit events form a chain so deletion or reordering changes the resulting hash. Export and deletion request are explicit, separate operations.",
            "Phase 1 does not claim HIPAA compliance, a business associate agreement, production identity proofing, real data, or a clinical safety case approved by a partner. Production work requires threat modeling, encryption and key management, tenant isolation, least privilege, retention rules, incident response, security testing, accessibility governance, clinical escalation ownership, and legal review.",
        ],
        "bullets": [
            "Human approval before any external action or activation.",
            "Professional resolution for medication conflicts and unidentified instructions.",
            "Minimum disclosure, expiry, revocation, and household boundary checks.",
            "Original-source retention, correction history, completion evidence, and audit receipts.",
        ],
    },
    {
        "title": "Evaluation evidence: controlled strengths and retained failures",
        "kicker": "Narrative page 7 of 15",
        "paragraphs": [
            "The original controlled evidence remains unchanged: Smart 40 passed 40/40 consecutive synthetic cases, and a focused benchmark passed 60/60 across responsibility fields, provenance and staleness, medication conflict, consent and privacy, capacity and delegation, and closed-loop workflow. These deterministic fixtures demonstrate implementation behavior, not clinical effectiveness or generalization.",
            "A separate external-style holdout of 20 cases was authored, ground-truthed, hashed, and locked before one execution. It passed 15/20. The five failures were retained without tuning or selected reruns: duplicate helper names, ambiguous numeric dates, clinical shorthand, concurrent conflicting corrections, and variable recurring exceptions. Two were safety-taxonomy failures because the system did not make escalation or conflict persistence explicit enough, although the harness did not execute a high-risk action.",
            "The repository now includes 18 unit/domain tests covering state transitions, completion evidence, superseded behavior, revoked and expired access, minimum disclosure, cross-household denial, prompt injection, outages, log redaction, correction history, audit integrity, export, and deletion request. Accessibility automation found no axe violations in three key screen states; manual caregiver screen-reader testing remains outstanding.",
        ],
        "table": (["Evidence", "Result", "Interpretation"], [
            ["Smart 40", "40/40", "Controlled consecutive fixtures"],
            ["Focused benchmark", "60/60", "Six deterministic behavior groups"],
            ["External-style holdout", "15/20", "Five failures retained; no rerun"],
            ["Caregiver outcomes", "Not measured", "No effect or effectiveness claim"],
        ], [2600, 1700, 5060]),
    },
    {
        "title": "Phase 2 pilot, milestones, and decision gates",
        "kicker": "Narrative page 8 of 15",
        "paragraphs": [
            "Phase 2 should be synthetic-data-first and caregiver-governed. Months 1–2 establish a signed scope, caregiver advisory group, accessible study materials, clinical/privacy/security owners, integration authority, incident process, and partner sandbox. Months 2–3 implement stable identity, consent and revocation, terminology-aware abstention, conflict-preserving corrections, locale confirmation, exception-aware scheduling, and partner-specific workflow adapters.",
            "Months 3–5 run iterative usability and burden studies with synthetic cases, then a narrowly scoped pilot only after governance approval. The primary burden measure is paired correct completion time across the fixed eight-task set. Secondary measures are interactions, help requests, corrections, completion, confidence, and effort. Guardrails are unsupported high-risk action, permission violation, unpreserved conflict, unverified completion, and unresolved accessibility blocker.",
            "A pilot advances only when high-risk holdout expectations pass, caregiver advisors approve authority language and privacy defaults, required accessibility checks pass, incident and escalation paths are tested, and the partner confirms authoritative sources and response events. Retained failures, adverse events, and subgroup differences remain visible. A faster workflow cannot offset a safety or permission failure.",
        ],
        "bullets": [
            "Go/no-go 1: real caregiver evidence and partner governance.",
            "Go/no-go 2: safety taxonomy, identity, consent, accessibility, and sandbox controls.",
            "Go/no-go 3: formative burden benefit without guardrail regression.",
            "Go/no-go 4: operational readiness and a sustainable, non-extractive support model.",
        ],
    },
    {
        "section": "Section 3: Usability and Integration",
        "title": "Caregiver-centered usability",
        "kicker": "Narrative page 9 of 15",
        "paragraphs": [
            "HEARTH uses a mission metaphor because caregivers need to answer a small set of operational questions quickly: What must happen next? What is unsafe or blocked? Who owns it? What evidence proves it happened? What may I share? The opening view is plain-language and consequence-first. Safety labels use words as well as color. Every screen repeats that the case is synthetic and not for clinical use.",
            "The guided reviewer path demonstrates nine state-changing controls in three to five minutes. A prominent Reset Reviewer Demo action restores the starting state. The evidence view distinguishes controlled tests, holdout failures, absent interview evidence, absent burden results, and simulated integrations. A separate timed study screen starts empty and records per-step time, interactions, help, corrections, confidence, effort, feedback, and a local JSON export.",
            "The application includes a skip link, labeled landmarks and controls, visible focus treatment, reduced-motion behavior, responsive reflow, keyboard-operable navigation, mobile navigation, form labels, and high-contrast tokens. Automated Playwright/axe checks cover welcome, demo, and burden-study states at desktop and 320 pixels. No axe violation was detected in those states, but automated testing is not a conformance claim.",
        ],
        "callout": ("Caregiver validation still required.", "Usability implementation evidence is stronger than usability outcome evidence. No caregiver task-success, satisfaction, screen-reader, or burden result is claimed."),
    },
    {
        "title": "Workflow integration and minimum necessary exchange",
        "kicker": "Narrative page 10 of 15",
        "paragraphs": [
            "HEARTH is designed to sit above existing documents and systems, not replace the legal record. The prototype uses ten synthetic source types: discharge PDF, current and prior medication lists, provider message, appointment instructions, home-health authorization, caregiver voice-note transcript, family availability, permission choices, and a FHIR-like sandbox bundle. Every extracted object links back to source and location.",
            "Phase 2 adapters may connect to FHIR or portal sandboxes, scheduling, messaging, pharmacy, home health, and aging-network resource directories, but each connector must declare authoritative fields, version and amendment behavior, response events, outage behavior, rate limits, and action authority. An adapter may prepare a request; it cannot promote “sent” to “completed” without an acknowledgement or outcome.",
            "Partner discovery identified public candidate organizations, including Illinois caregiver-resource and aging-network programs, but no outreach, commitment, data access, letter of intent, or integration agreement is claimed. A partner readiness checklist and integration-question set identify the work that must precede a pilot.",
        ],
        "table": (["Workflow", "Prepared disclosure", "Completion evidence"], [
            ["Transportation", "Date, time, location, mobility support, contact", "Helper accepts or task is reassigned"],
            ["Medication clarification", "Conflicting excerpts and callback under clinical purpose", "Authorized professional response + caregiver activation"],
            ["Home health", "Relevant instruction, equipment, training status", "Visit time and qualified owner confirmed"],
        ], [2100, 4200, 3060]),
    },
    {
        "title": "Adoption, burden, and equitable access",
        "kicker": "Narrative page 11 of 15",
        "paragraphs": [
            "The product should reduce coordination work without converting it into data-entry labor. The fixed burden-study protocol uses counterbalanced manual and assisted conditions on equivalent synthetic cases. Correct total completion time is primary; interactions, help, corrections, completion, confidence, and effort are secondary. Safety and permission errors are non-negotiable guardrails. At Phase 1 there are zero participants and no measured time-saving claim.",
            "Adoption begins with the transition workflow where responsibilities already exist, not with a demand that a household maintain a new record from scratch. Source import, review by exception, reusable permissions, printable missions, and task-specific helper views are intended to keep the burden low. Telephone, paper, interpreter, and navigator-supported pathways remain necessary for people with limited connectivity or digital comfort.",
            "Equitable implementation requires compensated caregiver participation, translation, accessible formats, disability accommodations, respite or scheduling support, low-bandwidth options, and subgroup review. A partner-supported free caregiver tier is preferred for core mission review, correction, permission, export, and deletion request. HEARTH will not use advertising or sell personal data.",
        ],
        "bullets": [
            "Measure who benefits and who must do more work.",
            "Report incomplete sessions, order effects, missing data, and disconfirming outcomes.",
            "Keep a human coordinator and non-digital path available.",
            "Do not scale until accessibility and burden guardrails are met.",
        ],
    },
    {
        "section": "Section 4: Alignment with Caregiver AI Principles",
        "title": "Human benefit, agency, transparency, and privacy",
        "kicker": "Narrative page 12 of 15",
        "paragraphs": [
            "Human benefit. HEARTH is designed around a caregiver’s operational burden and the care recipient’s ability to have preferences carried through. It does not optimize engagement or replace human relationships. A responsibility becomes executable only when its owner, dependencies, permission, evidence, and escalation are explicit.",
            "Agency and autonomy. The care recipient controls purpose-specific access and can revoke it. The caregiver approves external action and corrects interpretation. High-risk clinical ambiguity remains with qualified professionals. Supported decision-making is modeled as a daily responsibility, not treated as a one-time consent screen.",
            "Transparency and explainability. Each commitment shows source, excerpt, date, confidence, verification class, state, risk, completion criteria, and history. Each finding explains why a mission is blocked, what may continue, who can resolve it, and what evidence is required. Receipts show what was shared, withheld, approved, uncertain, and next.",
            "Privacy and data stewardship. Minimum necessary disclosure, expiry, revocation, household boundaries, log redaction, export, and deletion request are explicit controls. The production retention, identity, legal, and security design remains partner-defined and must be reviewed before real data.",
        ],
        "table": (["Principle", "Phase 1 evidence", "Remaining work"], [
            ["Human benefit", "Capacity, execution gaps, burden protocol", "Observed caregiver outcomes"],
            ["Agency", "Approval, correction, permissions, revocation", "Caregiver co-governance"],
            ["Transparency", "Provenance, findings, receipts, limitations", "Model monitoring"],
            ["Privacy", "Minimum disclosure and denial tests", "Production controls and legal review"],
        ], [1800, 4100, 3460]),
    },
    {
        "title": "Safety, fairness, accountability, and robustness",
        "kicker": "Narrative page 13 of 15",
        "paragraphs": [
            "Safety and reliability. H3/H4 boundaries, source verification, conflict detection, abstention, deterministic states, outage behavior, and completion evidence prevent common unsafe shortcuts. Protocol 9-Delta and medication-conflict examples demonstrate refusal to invent. The first-run holdout exposes where the taxonomy must improve before release.",
            "Fairness and inclusion. The current fixtures do not support a fairness-performance claim. Phase 2 will recruit across digital comfort, disability access needs, language, relationship, geography, work constraint, and care complexity; measure access and burden by subgroup; and allow non-digital pathways. A lack of observed subgroup harm in synthetic cases is not evidence of equity.",
            "Accountability. Humans remain named owners. Audit chains, correction histories, receipts, source versions, and separate export/deletion operations support review. Partner governance must assign clinical, privacy, security, accessibility, incident, and model-monitoring responsibility. HEARTH does not hide unresolved dependencies behind a generic “AI confidence” score.",
            "Robustness. Unit/domain tests, the Smart 40, focused 60, locked holdout, rendered-HTML tests, build/type checks, accessibility automation, and retained failures create a reproducible TRL-3 evidence base. Production robustness also requires real-world data validation, terminology services, identity, concurrency, localization, integration fault injection, and operational monitoring.",
        ],
        "callout": ("Honest boundary.", "The prototype demonstrates accountable design and controlled behavior. It does not establish clinical safety, fairness, real-world effectiveness, or regulatory compliance."),
    },
    {
        "section": "Section 5: Meritorious Prize Eligibility",
        "title": "TRL-3 accomplishment and meritorious use of prize support",
        "kicker": "Narrative page 14 of 15",
        "paragraphs": [
            "HEARTH is eligible as an early-stage Track 1 concept embodied in a functioning proof of concept. The repository contains a responsive product, synthetic fixtures, deterministic mission engine, evidence artifacts, tests, research and partner-readiness materials, affordability model, and submission documents. Reviewers can reset and execute the scenario without an account after public deployment.",
            "The work is meritorious because it does not use AI as a conversational veneer. It defines a new operational object and controlled state model for a caregiver problem that is routinely pushed onto families. The prototype demonstrates conflict preservation, minimum disclosure, capacity constraints, professional authority, audit receipts, and closed-loop outcomes in a coherent workflow.",
            "Prize support would be used to clear the evidence gaps rather than inflate the prototype: compensate a caregiver advisory group; conduct accessible formative research; implement the five holdout mitigations; complete production-grade identity, consent, privacy, security, and clinical governance; build one partner sandbox; and run a measured burden and usability evaluation.",
        ],
        "table": (["Six-month use", "Planning amount", "Evidence produced"], [
            ["Engineering and implementation", "$99,000", "Release-gated pilot build and sandbox"],
            ["Caregiver research/accessibility/translation", "$18,000", "Co-design, burden, usability, access results"],
            ["Clinical/privacy/security/evaluation advice", "$12,000", "Reviewed governance and safety case"],
            ["Cloud/tools and contingency", "$16,000", "Monitored pilot operations"],
        ], [3600, 1800, 3960]),
    },
    {
        "title": "Feasibility, sustainability, and closing case",
        "kicker": "Narrative page 15 of 15",
        "paragraphs": [
            "A directional six-month pilot budget is $145,000 for 100 enrolled households, or $1,450 per enrolled household including one-time setup, research, accessibility, governance, and contingency. A 1,000-active-household annual planning envelope is $380 per active household-year; a 10,000-household envelope is $170. These are transparent assumptions, not prices or validated forecasts.",
            "The proposed buyer is a health plan, provider, aging-network organization, employer benefit, or public program that benefits from safer transitions and reduced coordination friction. Core caregiver safety and control functions should not depend on an individual subscription. Revenue would come from implementation, support, and organization-level service—not advertising, sale of personal data, or referral commissions.",
            "HEARTH’s strongest Phase 1 contribution is a disciplined execution model with visible limits. It shows how AI can help caregivers without taking clinical authority, hiding uncertainty, over-sharing household information, or declaring victory when a message was merely sent. The next milestone is not a broader feature set; it is credible caregiver co-implementation and a governed partner pilot.",
        ],
        "bullets": [
            "Public reviewer prototype and backup guide require no login after deployment.",
            "Appendix contains architecture, evaluation, accessibility, research, partner, and cost evidence.",
            "Data Output Logs provide human-readable Smart 40, focused 60, and holdout records.",
            "Required pre-submission verification: contact fields, real caregiver evidence, partner status, and live public links.",
        ],
        "sources": True,
    },
]


def build_application():
    doc = Document()
    configure_document(doc, "HEARTH | Caregiver AI Challenge Phase 1 Application", total_pages=16)
    sec = doc.sections[0]
    sec.different_first_page_header_footer = False
    add_kicker(doc, "ACL Caregiver AI Challenge · Track 1")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run("HEARTH"), size=30, bold=True, color=PINE, font="Georgia")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(25)
    set_run(p.add_run("Care information, made executable."), size=15, italic=True, color=PINE_2)
    add_callout(doc, "Phase 1 submission.", "A caregiver-first, permission-aware, closed-loop coordination proof of concept. All product records and people are synthetic. Not for clinical use.", PALE)
    doc.add_paragraph()
    add_table(doc, ["Submission field", "Value"], [
        ["Challenge track", "Track 1 · Early-stage / TRL 3"],
        ["Applicant / team", "HEARTH · Shobhit Kapoor"],
        ["Team lead", "Shobhit Kapoor"],
        ["Contact email", "[COMPLETE BEFORE SUBMISSION — contact email not supplied]"],
        ["Contact phone", "[COMPLETE BEFORE SUBMISSION — phone not supplied]"],
        ["Organization / status", "[COMPLETE BEFORE SUBMISSION — legal/organizational status not supplied]"],
        ["Public prototype", "HEARTH public reviewer site · descriptive link on narrative page 15"],
        ["Submission date", "29 July 2026"],
    ], [2800, 6560], font_size=10)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Submission-ready except for explicitly marked applicant fields and missing real caregiver evidence."), size=10, bold=True, color=RED)
    page_break(doc)

    for index, page in enumerate(APPLICATION_PAGES):
        if page.get("section"):
            doc.add_paragraph(page["section"], style="Heading 1")
        add_kicker(doc, page["kicker"])
        doc.add_paragraph(page["title"], style="Heading 2")
        for text in page.get("paragraphs", []):
            add_body(doc, text)
        if page.get("bullets"):
            add_bullets(doc, page["bullets"], compact=True)
        if page.get("table"):
            headers, rows, widths = page["table"]
            add_table(doc, headers, rows, widths, font_size=9)
        if page.get("callout"):
            add_callout(doc, page["callout"][0], page["callout"][1])
        if page.get("sources"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            set_run(p.add_run("Official challenge sources: "), size=9, bold=True)
            for idx, (label, url) in enumerate([
                ("Public reviewer prototype", "https://hearth-care-execution.shobhit1kapoor.chatgpt.site"),
                ("Backup reviewer guide", "https://hearth-care-execution.shobhit1kapoor.chatgpt.site/reviewer-demo.html"),
                ("Challenge page", "https://acl.gov/caregiver-ai-challenge"),
                ("Track 1 criteria", "https://acl.gov/caregiver-ai-judging-track1"),
                ("Application outline", "https://acl.gov/caregiver-ai-application-outline"),
                ("Technology readiness guide", "https://acl.gov/caregiver-ai-tech-readiness-guide"),
                ("Definitions and FAQ", "https://acl.gov/caregiver-ai-definitions-faq"),
            ]):
                if idx:
                    set_run(p.add_run(" · "), size=9, color=MUTED)
                add_hyperlink(p, label, url)
        if index < len(APPLICATION_PAGES) - 1:
            page_break(doc)
    doc.core_properties.title = "HEARTH Phase 1 Application"
    doc.core_properties.subject = "ACL Caregiver AI Challenge Track 1"
    doc.core_properties.keywords = "caregiver, AI, care coordination, TRL 3"
    path = OUT / "HEARTH_Phase1_Application.docx"
    doc.save(path)
    return path


APPENDIX_PAGES = [
    ("Appendix A — Compliance and judging map", [
        "This appendix is limited to ten pages and does not extend the narrative. The application uses a one-page cover and fifteen narrative pages with the exact official section headings.",
        "Compliance status is explicit: one-inch margins, 11-point body, 9-point tables, accessible Word/PDF output, no government seal, and descriptive public links. Required applicant contact fields remain visibly marked because they were not supplied.",
    ], (["Criterion", "Primary application pages", "Evidence"], [
        ["Need and solution", "1–4", "Synthetic case, CCO design, caregiver-evidence gate"],
        ["Implementation", "5–8", "Architecture, safety, evaluation, pilot gates"],
        ["Usability/integration", "9–11", "Responsive product, a11y, workflow adapters"],
        ["AI principles", "12–13", "Agency, transparency, privacy, safety, fairness"],
        ["Merit/eligibility", "14–15", "TRL-3 artifacts, prize use, sustainability"],
    ], [2500, 2700, 4160])),
    ("Appendix B — System architecture and trust boundaries", [
        "Source artifacts enter an immutable source store. Candidate interpretation is separated from deterministic mission controls. Commitments, permissions, audit events, and workflow adapters remain distinct. An adapter cannot change authority merely because it can reach an external system.",
        "Trust boundaries include household tenancy, purpose-specific disclosure, high-risk professional review, caregiver approval, external acknowledgement, and completion evidence. Outages preserve open work rather than producing a success state.",
    ], (["Boundary", "Control", "Phase 2 verification"], [
        ["Source → interpretation", "Provenance and injection quarantine", "Adversarial document tests"],
        ["Interpretation → activation", "Safety level + human approval", "Clinical governance sign-off"],
        ["Household → helper", "Purpose, fields, expiry, revocation", "Identity and consent testing"],
        ["HEARTH → external system", "Prepare/approve/send/outcome states", "Sandbox fault injection"],
    ], [2600, 3500, 3260])),
    ("Appendix C — Care commitment object and lifecycle", [
        "The 26 synthetic care commitment objects include responsibility, owner, alternates, source, location, excerpt, date, verification class, confidence, due window, dependencies, equipment, skill, privacy, consent, risk, safety level, approval, completion criteria, evidence, backup, escalation, state, duration, and history.",
        "Valid transitions are explicit. Completion and verification require evidence. Superseded objects cannot reactivate. A sent request ordinarily moves to Awaiting external response, not Completed.",
    ], (["Lifecycle example", "Allowed next state", "Required evidence"], [
        ["Awaiting acceptance", "Accepted or reassigned", "Helper response"],
        ["In progress", "Awaiting external response", "Request sent + recipient"],
        ["Awaiting external response", "Completed", "Acknowledgement/outcome"],
        ["Completed", "Verified", "Reviewer verification"],
        ["Superseded", "None", "Immutable history only"],
    ], [2800, 2800, 3760])),
    ("Appendix D — Controlled validation", [
        "Smart 40 and the focused 60 are deterministic, consecutive, synthetic checks. They are valuable for reproducibility but intentionally narrow. They do not justify a real-world precision, recall, fairness, clinical, or caregiver-outcome claim.",
        "The suite retains full per-case input, expected behavior, actual output, pass state, model/configuration, timestamp, human-review flag, and correction field in the Data Output Logs.",
    ], (["Suite", "Cases", "Pass", "Key boundary"], [
        ["Smart 40", "40", "40", "Messy input, injection, uncertainty, operational failure"],
        ["Focused benchmark", "60", "60", "Provenance, conflict, privacy, capacity, lifecycle"],
        ["Required abstention examples", "4", "4", "No unsupported high-risk action"],
    ], [2900, 1500, 1500, 3460])),
    ("Appendix E — Locked external-style holdout", [
        f"The 20 cases and ground truth were hashed before the first execution. The run passed {json.loads((ROOT/'evidence/holdout/results.json').read_text())['summary']['passed']}/20 and retained five failures with no selected rerun.",
        "Mitigations are stable person identifiers and disambiguation, locale confirmation for numeric dates, terminology-aware escalation, conflict-preserving concurrent correction, and exception-aware recurring schedules. The latter three are release gates when safety or action semantics are affected.",
    ], (["Failure", "Observed behavior", "Release implication"], [
        ["Duplicate name", "First-match assignment", "Require stable identity"],
        ["Ambiguous date", "US-locale parse", "Require locale confirmation"],
        ["Clinical shorthand", "Generic needs review", "Make professional escalation explicit"],
        ["Conflicting corrections", "Last write wins", "Persist unresolved conflict"],
        ["Recurring exception", "Daily schedule only", "Model exception rule"],
    ], [2600, 3100, 3660])),
    ("Appendix F — Timed burden study", [
        "The browser contains an empty, resettable study mode for manual and HEARTH-assisted conditions. It records eight fixed tasks, per-step and total time, interactions, help, corrections, confidence, effort, optional de-identified feedback, and a local JSON export.",
        "No participant record is preloaded; no result is transmitted. The proposed design is within-participant and counterbalanced. Correct completion time is primary, with safety and permission errors as guardrails.",
    ], (["Measure", "Phase 1 value", "Phase 2 analysis"], [
        ["Participants", "0", "5–8 formative; larger pilot as governed"],
        ["Time saving", "Not measured", "Median paired difference + range"],
        ["Interactions/help/corrections", "Instrument ready", "Paired and task-level analysis"],
        ["Confidence/effort", "Instrument ready", "Distribution and disconfirming cases"],
    ], [3000, 2500, 3860])),
    ("Appendix G — Safety, privacy, and security", [
        "The safety model separates H0 administrative support, H1 reviewable logistics, H2 sensitive/high-consequence preparation, H3 professional review, and H4 emergency boundaries. Conflicts, uncertainty, and stale sources remain visible.",
        "Privacy controls deny cross-household, revoked, expired, or purposeless access and disclose only allowed fields. Security and compliance remain pre-production work; the appendix does not imply certification.",
    ], (["Threat", "Phase 1 control", "Production gate"], [
        ["Prompt injection in source", "Quarantine as untrusted text", "Adversarial monitoring"],
        ["Over-disclosure", "Purpose-specific allowlist", "Policy engine + identity"],
        ["Silent source overwrite", "Versions and correction history", "Immutable store"],
        ["False completion", "Outcome evidence requirement", "Adapter acknowledgements"],
        ["Audit tampering", "Event hash chain", "Signed/tamper-evident log"],
    ], [2600, 3000, 3760])),
    ("Appendix H — Accessibility evidence", [
        "The automated production-build audit covers three key states with Playwright and axe-core, keyboard skip behavior, sampled focus visibility, 320-pixel reflow, mobile navigation, and a token contrast matrix. The final recorded run found zero axe violations in the audited states.",
        "This does not establish conformance or caregiver usability. Screen-reader testing with caregivers, assistive-technology diversity, document reading order in all viewers, and partner accessibility governance remain Phase 2 tasks.",
    ], (["Check", "Result", "Caveat"], [
        ["Axe: welcome/demo/study", "0 violations", "Three states only"],
        ["Keyboard skip", "Pass", "Automated"],
        ["320px overflow", "Pass", "Audited states"],
        ["Token contrast", "AA pairs pass after correction", "Sampled tokens"],
        ["Caregiver screen-reader test", "Not performed", "Required"],
    ], [3100, 2100, 4160])),
    ("Appendix I — Caregiver and partner evidence gates", [
        "No consented caregiver research record was supplied. Therefore there is no participant register, empirical theme, authentic quote, caregiver-attributed design change, or measured burden result. The research folder contains transparent empty-state artifacts and a co-implementation plan.",
        "No partner has been contacted or committed through this build. The repository contains a public candidate list, nonbinding letter template, readiness checklist, integration questions, and unresolved-dependency register. Candidates are not represented as relationships.",
    ], (["Evidence class", "Current status", "Action before claim"], [
        ["Caregiver interview", "Missing", "Consent + de-identified source + verified analysis"],
        ["Caregiver quote", "Missing", "Source check + quote permission"],
        ["Partner commitment", "None", "Authorized outreach + written evidence"],
        ["Integration", "Simulation", "Sandbox, authority, governance, testing"],
    ], [3000, 2500, 3860])),
    ("Appendix J — Cost, sustainability, and roadmap", [
        "The six-month pilot estimate is $145,000 for 100 enrolled households. It includes engineering, implementation, compensated caregiver research, accessibility and translation, clinical/privacy/security advice, cloud/tools, and contingency.",
        "The sustainability strategy protects core caregiver control functions and avoids advertising or sale of personal data. Phase 2 tests buyer value, partner burden, caregiver benefit, equity, and actual unit cost before a pricing commitment.",
    ], (["Milestone", "Timing", "Decision evidence"], [
        ["Governance + partner scope", "Months 1–2", "Signed roles, accessible protocol"],
        ["Safety and platform hardening", "Months 2–3", "Holdout gates + security review"],
        ["Formative evaluation", "Months 3–5", "Usability, burden, guardrails"],
        ["Pilot decision", "Month 6", "Caregiver/partner go-no-go"],
    ], [3000, 1800, 4560])),
]


def build_appendix():
    doc = Document()
    configure_document(doc, "HEARTH | Phase 1 Appendix", total_pages=10)
    for idx, (title, paragraphs, table) in enumerate(APPENDIX_PAGES):
        add_kicker(doc, f"Phase 1 appendix · page {idx + 1} of 10")
        doc.add_paragraph(title, style="Heading 1")
        for paragraph in paragraphs:
            add_body(doc, paragraph)
        headers, rows, widths = table
        add_table(doc, headers, rows, widths, font_size=9)
        if idx == 0:
            add_callout(doc, "Boundary.", "Appendix evidence does not cure the missing real caregiver record or partner commitment.")
        if idx < len(APPENDIX_PAGES) - 1:
            page_break(doc)
    doc.core_properties.title = "HEARTH Phase 1 Appendix"
    path = OUT / "HEARTH_Phase1_Appendix.docx"
    doc.save(path)
    return path


def build_logs():
    validation = json.loads((ROOT / "evidence/validation-results.json").read_text(encoding="utf-8"))
    holdout = json.loads((ROOT / "evidence/holdout/results.json").read_text(encoding="utf-8"))
    doc = Document()
    configure_document(doc, "HEARTH | Human-readable Data Output Logs")
    add_kicker(doc, "Phase 1 evaluation evidence")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    set_run(p.add_run("HEARTH Data Output Logs"), size=27, bold=True, color=PINE, font="Georgia")
    add_body(doc, "Human-readable controlled validation and first-run holdout records")
    add_callout(doc, "Interpretation boundary.", "These logs describe deterministic TRL-3 behavior on synthetic or external-style cases. They are not clinical evidence, caregiver-outcome evidence, or a claim of generalization.")
    add_table(doc, ["Suite", "Total", "Passed", "Failed"], [
        ["Smart 40", "40", "40", "0"],
        ["Focused benchmark", "60", "60", "0"],
        ["External-style holdout", str(holdout["summary"]["total"]), str(holdout["summary"]["passed"]), str(holdout["summary"]["failed"])],
    ], [3900, 1800, 1800, 1860], font_size=10)
    page_break(doc)

    def compact_line(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        set_run(p.add_run(label + ": "), size=11, bold=True, color=PINE)
        set_run(p.add_run(str(value)), size=11)

    def case_heading(text, passed):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), MINT if passed else "F7E7E3")
        p_pr.append(shd)
        set_run(p.add_run(text), size=11, bold=True, color=PINE if passed else RED)

    doc.add_paragraph("Controlled validation — Smart 40 and focused 60", style="Heading 1")
    add_body(doc, f"Configuration: {validation['run']['configuration']}. Consecutive: {validation['run']['consecutive']}. Selected reruns: {validation['run']['selectedReruns']}.")
    for index, item in enumerate(validation["results"]):
        case_heading(f"{item['id']} · {item['group']} · {'PASS' if item['passed'] else 'FAIL'}", item["passed"])
        compact_line("Input", item["input"])
        compact_line("Expected", item["expected"])
        compact_line("Actual", item["actual"])
        compact_line("Review", "Human review required" if item["humanReviewRequired"] else "No additional human-review flag")
        if (index + 1) % 4 == 0 and index + 1 < len(validation["results"]):
            page_break(doc)
            doc.add_paragraph("Controlled validation — continued", style="Heading 2")

    page_break(doc)
    doc.add_paragraph("External-style holdout — first and only recorded run", style="Heading 1")
    add_body(doc, f"Executed {holdout['run']['executedAt']}. Cases SHA-256 {holdout['run']['casesSha256']}. Ground truth SHA-256 {holdout['run']['groundTruthSha256']}. No selected reruns.")
    for index, item in enumerate(holdout["results"]):
        status = "PASS" if item["passed"] else "FAIL — retained"
        if item["safetyCritical"] and not item["passed"]:
            status += " · safety-critical"
        case_heading(f"{item['id']} · {item['kind']} · {status}", item["passed"])
        compact_line("Input", item["input"])
        compact_line("Expected", item["expected"])
        compact_line("Observed", item["actual"])
        compact_line("Behavior", item["behavior"])
        compact_line("Correction", item["correction"])
        if (index + 1) % 3 == 0 and index + 1 < len(holdout["results"]):
            page_break(doc)
            doc.add_paragraph("External-style holdout — continued", style="Heading 2")

    page_break(doc)
    doc.add_paragraph("Required abstention output", style="Heading 1")
    add_callout(doc, "Protocol 9-Delta.", "I cannot identify Protocol 9-Delta as a verified instruction. I will not invent its meaning or apply it. Please provide an approved source or request review from a qualified professional.", PALE)
    add_body(doc, "This exact output is retained in Smart 40 case S40-S01 and is reproduced here for reviewer readability.")
    doc.add_paragraph("Known limitations carried forward", style="Heading 2")
    add_bullets(doc, [
        "Duplicate-name assignment needs stable identity and disambiguation.",
        "Ambiguous numeric dates need locale confirmation.",
        "Clinical shorthand needs an explicit terminology-aware professional escalation.",
        "Concurrent corrections need conflict-preserving merge behavior.",
        "Variable recurring exceptions need an exception-aware schedule.",
        "No caregiver outcome, clinical effectiveness, fairness, or burden-reduction result is contained in these logs.",
    ])
    doc.core_properties.title = "HEARTH Data Output Logs"
    path = OUT / "HEARTH_Data_Output_Logs.docx"
    doc.save(path)
    return path


def build_markdown_summaries():
    plain = """# HEARTH plain-language summary

HEARTH helps a family caregiver turn scattered care information into a plan that can actually be carried out. It identifies what needs to happen, who owns each responsibility, what information supports it, what could block it, what may be shared, and what evidence shows the work is finished.

HEARTH does not diagnose, prescribe, choose between conflicting medication instructions, or act without permission. When an instruction is unclear or requires professional authority, it stops and asks for an approved source or qualified review.

The Phase 1 prototype uses an entirely synthetic household. Ten synthetic sources become 26 responsibilities. Reviewers can resolve nine controlled gaps and watch the mission move from “Not executable” to “Ready with controls.”

Evidence includes 40/40 controlled cases, a 60/60 focused benchmark, and a separately locked 15/20 external-style holdout with five failures retained. The product also contains an empty timed burden-study instrument. No real caregiver interview evidence, measured time saving, clinical effectiveness, partner commitment, or live health-system integration is claimed.

The next phase is to co-implement HEARTH with compensated caregivers and a governed partner, fix the retained holdout failures, complete production safety/privacy/accessibility controls, and measure whether the workflow genuinely reduces burden.
"""
    executive = """# HEARTH executive summary

HEARTH is a caregiver-first care-execution assurance layer for the operational gap after discharge. It compiles fragmented documents and messages into source-grounded care commitment objects with owners, safe windows, dependencies, permissions, completion evidence, backups, escalation, and lifecycle state.

The functioning TRL-3 browser prototype uses 10 synthetic sources and 26 commitments. It demonstrates conflict preservation, high-risk abstention, minimum-necessary disclosure, caregiver-capacity constraints, closed-loop external outcomes, correction history, and accountability receipts. All external outcomes are labeled simulations.

Controlled evidence is reproducible: Smart 40 passed 40/40 and the focused benchmark passed 60/60. A separately authored and pre-locked external-style holdout passed 15/20 on its first run; five failures were retained and two safety-taxonomy gaps remain Phase 2 release blockers. Automated Playwright/axe testing found zero violations in three key states. The burden-study instrument is implemented but has zero participants and no time-saving claim.

The most important missing evidence is real caregiver input. No consented interview record was supplied, so the application claims no participants, quotes, empirical themes, or caregiver-attributed design changes. No partner commitment is claimed. Phase 2 prioritizes compensated caregiver co-governance, a partner sandbox, the five holdout mitigations, production identity/consent/security controls, and a counterbalanced burden and usability evaluation.
"""
    (OUT / "HEARTH_Plain_Language_Summary.md").write_text(plain, encoding="utf-8")
    (OUT / "HEARTH_Executive_Summary.md").write_text(executive, encoding="utf-8")


paths = [build_application(), build_appendix(), build_logs()]
build_markdown_summaries()
for path in paths:
    print(path)
